import os

import docx
from docx.shared import Inches as DocxInches, Pt as DocxPt
from docx.shared import RGBColor as DocxRGBColor
from typing import Optional, Dict, Any, List

UPLOAD_FOLDER_DOCX = 'documents_fastmcp' # New folder for Word documents
if not os.path.exists(UPLOAD_FOLDER_DOCX):
    os.makedirs(UPLOAD_FOLDER_DOCX)


class WordDocumentEditor:
    """
    Handles Word document (.docx) creation and modification using python-docx.
    """
    def __init__(self):
        self.doc = None
        self.current_filename = None

    def _ensure_document_loaded(self):
        if self.doc is None:
            raise ValueError("No document loaded. Call create_docx_document() or load_docx_document() first.")

    def create_document(self, filename: str = "new_document.docx") -> str:
        """Creates a new, blank Word document in memory."""
        self.doc = docx.Document()
        if os.path.isabs(filename):
            self.current_filename = filename
        else:
            self.current_filename = os.path.join(UPLOAD_FOLDER_DOCX, filename)
        # Add an initial empty paragraph to ensure the document is not completely empty
        if not self.doc.paragraphs:
            self.doc.add_paragraph("")
        return f"New Word document '{self.current_filename}' created and ready in memory."

    def load_document(self, filename: str) -> str:
        """Loads an existing Word document from the 'documents_fastmcp' directory."""
        filepath = os.path.join(UPLOAD_FOLDER_DOCX, filename) if not os.path.isabs(filename) else filename
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Word document file '{filepath}' not found.")
        self.doc = docx.Document(filepath)
        self.current_filename = filepath
        return f"Word document '{self.current_filename}' loaded successfully."

    def save_document(self, filename: Optional[str] = None) -> str:
        """Saves the current in-memory Word document."""
        self._ensure_document_loaded()
        save_path = self.current_filename
        if filename:
            save_path = os.path.join(UPLOAD_FOLDER_DOCX, filename) if not os.path.isabs(filename) else filename
            self.current_filename = save_path
        if not save_path:
            raise ValueError("Filename not specified for saving Word document.")
        save_dir = os.path.dirname(save_path)
        if save_dir and not os.path.exists(save_dir):
            os.makedirs(save_dir)
        self.doc.save(save_path)
        return f"Word document saved to '{save_path}'."

    def add_paragraph(self, text: str, style: Optional[str] = None) -> Dict[str, str]:
        """Adds a paragraph with optional style (e.g., 'Normal', 'BodyText', 'Heading1')."""
        self._ensure_document_loaded()
        paragraph = self.doc.add_paragraph(str(text), style=style)
        return {"text_added": paragraph.text, "style_applied": style if style else "Normal (default)"}

    def add_heading(self, text: str, level: int = 1) -> Dict[str, Any]:
        """Adds a heading with a specific level (0-9). Level 0 is Title."""
        self._ensure_document_loaded()
        if not (0 <= level <= 9):
            raise ValueError("Heading level must be between 0 and 9.")
        heading = self.doc.add_heading(str(text), level=level)
        return {"text_added": heading.text, "level": level}

    def add_styled_text_to_paragraph(self, text_runs: List[Dict[str, Any]], paragraph_style: Optional[str] = None) -> str:
        """
        Adds a paragraph composed of multiple text runs with individual styling.
        Each run in text_runs is a dict: {"text": "str", "bold": bool, "italic": bool, 
                                          "font_size_pt": int, "font_name": "str", "font_color_rgb": [R,G,B]}
        """
        self._ensure_document_loaded()
        p = self.doc.add_paragraph(style=paragraph_style)
        full_text = []
        for run_info in text_runs:
            text = str(run_info.get("text", ""))
            run = p.add_run(text)
            full_text.append(text)
            if run_info.get("bold"): run.bold = True
            if run_info.get("italic"): run.italic = True
            if "font_size_pt" in run_info: run.font.size = DocxPt(int(run_info["font_size_pt"])) # Use aliased DocxPt
            if "font_name" in run_info: run.font.name = str(run_info["font_name"])
            if "font_color_rgb" in run_info:
                color = run_info["font_color_rgb"]
                if isinstance(color, list) and len(color) == 3:
                    run.font.color.rgb = DocxRGBColor(*color) # Use aliased DocxRGBColor
                else:
                    print(f"Warning: Invalid font_color_rgb format for run '{text}'. Expected [R,G,B].")
        return "".join(full_text)


    def add_table(self, rows: int, cols: int, data_list: Optional[List[List[str]]] = None, style: Optional[str] = 'TableGrid') -> Dict[str, Any]:
        """Adds a table with specified rows, columns, optional data, and style."""
        self._ensure_document_loaded()
        if rows <= 0 or cols <= 0:
            raise ValueError("Number of rows and columns must be positive.")
        table = self.doc.add_table(rows=rows, cols=cols, style=style)
        if data_list:
            if len(data_list) != rows or not all(len(row_data) == cols for row_data in data_list):
                raise ValueError("Data_list dimensions must match specified rows and columns.")
            for r_idx, row_content in enumerate(data_list):
                for c_idx, cell_content in enumerate(row_content):
                    table.cell(r_idx, c_idx).text = str(cell_content)
        return {"rows": rows, "cols": cols, "style": style, "data_populated": bool(data_list)}

    def add_picture(self, image_path: str, width_inch: Optional[float] = None, height_inch: Optional[float] = None) -> str:
        """Adds a picture from a local path, optionally scaled."""
        self._ensure_document_loaded()
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Image file not found at '{image_path}'.")

        width_val = DocxInches(width_inch) if width_inch is not None else None
        height_val = DocxInches(height_inch) if height_inch is not None else None

        try:
            if width_val and height_val:
                 self.doc.add_picture(image_path, width=width_val, height=height_val)
            elif width_val:
                 self.doc.add_picture(image_path, width=width_val)
            elif height_val: 
                 self.doc.add_picture(image_path, height=height_val)
            else:
                 self.doc.add_picture(image_path)
            return f"Picture added from '{os.path.basename(image_path)}'."
        except Exception as e:
            raise ValueError(f"Could not add picture from '{image_path}': {e}")


    def add_page_break(self) -> str:
        """Adds a manual page break."""
        self._ensure_document_loaded()
        self.doc.add_page_break()
        return "Page break added."
